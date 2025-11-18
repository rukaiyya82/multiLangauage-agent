<<<<<<< HEAD
"""
AI Multi-Language Corporate Communicator Agent
Fixed version: works with Google GenAI SDK and browser TTS
"""

import os
import io
import base64
from typing import Optional, List

import streamlit as st
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader

# Load env
from dotenv import load_dotenv
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Optional translator
try:
    from googletrans import Translator
    _translator = Translator()
except Exception:
    _translator = None

# Optional TTS fallback
try:
    from gtts import gTTS
except Exception:
    gTTS = None

# Optional microphone capture
try:
    from streamlit_webrtc import webrtc_streamer, WebRtcMode, ClientSettings
    webrtc_available = True
except Exception:
    webrtc_available = False

# -----------------------------
# GenAI SDK
# -----------------------------
try:
    from google.genai import Client
    genai_client = Client(api_key=api_key)
except Exception as e:
    st.error(f"GenAI Client not initialized: {e}")
    genai_client = None

# -----------------------------
# Config & constants
# -----------------------------
APP_TITLE = "LangBot: An AI-Powered Multi-Language Translation Assistant"
TONE_OPTIONS = ["Professional", "Neutral", "Friendly", "Assertive", "Concise", "Formal", "Casual"]

CORPORATE_TEMPLATES = {
    "General Professional Email": "Dear {recipient},\n\n{body}\n\nBest regards,\n{sender}",
    "Meeting Request": "Subject: Meeting Request - {subject}\n\nHi {recipient},\n\nI would like to request a meeting regarding {subject}. Proposed times: {times}.\n\nBest regards,\n{sender}"
}

LANGUAGES = {
    "English": "en",
    "Hindi": "hi",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Chinese (Simplified)": "zh-cn",
    "Japanese": "ja",
    "Korean": "ko",
    "Arabic": "ar",
    "Portuguese": "pt",
    "Russian": "ru",
}

# -----------------------------
# Helper functions
# -----------------------------
def read_uploaded_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith('.txt'):
        return uploaded_file.read().decode('utf-8', errors='ignore')
    if name.endswith('.docx'):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith('.pdf'):
        try:
            reader = PdfReader(uploaded_file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except:
            return ""
    return uploaded_file.read().decode('utf-8', errors='ignore')

def translate_text(text: str, dest_lang: str) -> str:
    if not _translator:
        return text
    try:
        res = _translator.translate(text, dest=dest_lang)
        return res.text
    except Exception:
        return text

def build_prompt(user_text: str, input_lang: str, output_lang: str, tone: str,
                 template_name: Optional[str], template_fields: dict, history: List[str]) -> str:
    history_part = ''
    if history:
        history_part = "\n---\nPrevious conversation:\n" + "\n".join(history[-10:]) + "\n---\n"
    if template_name and template_name in CORPORATE_TEMPLATES:
        template = CORPORATE_TEMPLATES[template_name].format(**template_fields)
        body = template.replace('{body}', user_text)
    else:
        body = user_text
    prompt_parts = [
        "You are a senior corporate communications assistant.",
        f"Input language: {input_lang}. Output language: {output_lang}.",
        f"Tone: {tone}.",
        "Produce a polished, professional message suitable for corporate use.",
        history_part,
        "Message:",
        body,
        "-- End of instructions --"
    ]
    return "\n".join([p for p in prompt_parts if p])

def generate_docx_from_text(text: str) -> bytes:
    doc = Document()
    for line in text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def make_eml(subject: str, body: str, sender: str='sender@example.com', recipient: str='recipient@example.com') -> bytes:
    eml = f"""From: {sender}
To: {recipient}
Subject: {subject}
MIME-Version: 1.0
Content-Type: text/plain; charset=utf-8

{body}"""
    return eml.encode('utf-8')

def serve_gtts_to_user(text: str) -> Optional[bytes]:
    if not gTTS:
        return None
    t = gTTS(text)
    bio = io.BytesIO()
    t.write_to_fp(bio)
    bio.seek(0)
    return bio.read()

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title=APP_TITLE, layout='centered')

if 'history' not in st.session_state:
    st.session_state.history = []

if 'last_generated' not in st.session_state:
    st.session_state.last_generated = ''

st.title(APP_TITLE)

input_lang_name = st.selectbox('Input language', list(LANGUAGES.keys()))
output_lang_name = st.selectbox('Output language', list(LANGUAGES.keys()))
tone = st.selectbox('Tone', TONE_OPTIONS)

uploaded = st.file_uploader('Upload source file', type=['txt','docx','pdf'])
uploaded_text = read_uploaded_file(uploaded) if uploaded else ''
user_text = st.text_area('Source text', value=uploaded_text, height=200)

template_choice = st.selectbox('Template', ['(none)'] + list(CORPORATE_TEMPLATES.keys()))
template_fields = {}
if template_choice != '(none)':
    template_fields['recipient'] = st.text_input('Recipient', 'Team')
    template_fields['sender'] = st.text_input('Sender', 'Your Name')
    template_fields['subject'] = st.text_input('Subject', 'Updates')
    template_fields['times'] = st.text_input('Proposed times', 'Mon/Tue 10-11am')

generate_button = st.button('Generate Message')
reset_button = st.button('Reset Memory')

if reset_button:
    st.session_state.history = []
    st.session_state.last_generated = ''
    st.success('Memory reset.')

if generate_button and genai_client:
    # Translate input to English for processing
    internal_input = user_text
    if input_lang_name != 'English' and _translator:
        internal_input = translate_text(user_text, 'en')

    prompt = build_prompt(internal_input, input_lang_name, output_lang_name, tone,
                          (template_choice if template_choice != '(none)' else None),
                          template_fields, st.session_state.history)

    # Generate with GenAI
    try:
        response = genai_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        result = response.candidates[0].content.parts[0].text
    except Exception as e:
        st.error(f"Generation failed: {e}")
        result = ''

    # Translate back to output language
    final_output = result
    if output_lang_name != 'English' and _translator:
        final_output = translate_text(result, LANGUAGES[output_lang_name])

    st.session_state.last_generated = final_output
    st.session_state.history.append(f"User request: {user_text}")
    st.session_state.history.append(f"Generated: {final_output}")
    st.success('Message generated.')

# Preview & downloads
if st.session_state.last_generated:
    escaped_text = st.session_state.last_generated.replace("`","\\`").replace("\n","\\n")
    tts_html = f"""
    <div>
    <button id='speak' style='background:#b8860b; border:none; padding:8px 12px; border-radius:6px; font-weight:700; color:#071018'>Speak</button>
    <script>
    const text = `{escaped_text}`;
    const btn = document.getElementById('speak');
    btn.onclick = () => {{
        const utter = new SpeechSynthesisUtterance(text);
        utter.lang='{LANGUAGES[output_lang_name]}';
        window.speechSynthesis.speak(utter);
    }};
    </script>
    </div>
    """
    st.markdown(tts_html, unsafe_allow_html=True)
    st.text_area('Generated Message', st.session_state.last_generated, height=200)
    st.download_button('Download TXT', st.session_state.last_generated, 'message.txt')
    st.download_button('Download DOCX', generate_docx_from_text(st.session_state.last_generated), 'message.docx')

st.sidebar.title('Session Memory')
if st.session_state.history:
    for i, item in enumerate(st.session_state.history[::-1]):
        st.sidebar.write(f"{len(st.session_state.history)-i}. {item[:200]}{'...' if len(item)>200 else ''}")
else:
    st.sidebar.write("No history yet.")
=======
"""
AI Multi-Language Corporate Communicator Agent
Fixed version: works with Google GenAI SDK and browser TTS
"""

import os
import io
import base64
from typing import Optional, List

import streamlit as st
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader

# Load env
from dotenv import load_dotenv
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Optional translator
try:
    from googletrans import Translator
    _translator = Translator()
except Exception:
    _translator = None

# Optional TTS fallback
try:
    from gtts import gTTS
except Exception:
    gTTS = None

# Optional microphone capture
try:
    from streamlit_webrtc import webrtc_streamer, WebRtcMode, ClientSettings
    webrtc_available = True
except Exception:
    webrtc_available = False

# -----------------------------
# GenAI SDK
# -----------------------------
try:
    from google.genai import Client
    genai_client = Client(api_key=api_key)
except Exception as e:
    st.error(f"GenAI Client not initialized: {e}")
    genai_client = None

# -----------------------------
# Config & constants
# -----------------------------
APP_TITLE = "LangBot: An AI-Powered Multi-Language Translation Assistant"
TONE_OPTIONS = ["Professional", "Neutral", "Friendly", "Assertive", "Concise", "Formal", "Casual"]

CORPORATE_TEMPLATES = {
    "General Professional Email": "Dear {recipient},\n\n{body}\n\nBest regards,\n{sender}",
    "Meeting Request": "Subject: Meeting Request - {subject}\n\nHi {recipient},\n\nI would like to request a meeting regarding {subject}. Proposed times: {times}.\n\nBest regards,\n{sender}"
}

LANGUAGES = {
    "English": "en",
    "Hindi": "hi",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Chinese (Simplified)": "zh-cn",
    "Japanese": "ja",
    "Korean": "ko",
    "Arabic": "ar",
    "Portuguese": "pt",
    "Russian": "ru",
}

# -----------------------------
# Helper functions
# -----------------------------
def read_uploaded_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith('.txt'):
        return uploaded_file.read().decode('utf-8', errors='ignore')
    if name.endswith('.docx'):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith('.pdf'):
        try:
            reader = PdfReader(uploaded_file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except:
            return ""
    return uploaded_file.read().decode('utf-8', errors='ignore')

def translate_text(text: str, dest_lang: str) -> str:
    if not _translator:
        return text
    try:
        res = _translator.translate(text, dest=dest_lang)
        return res.text
    except Exception:
        return text

def build_prompt(user_text: str, input_lang: str, output_lang: str, tone: str,
                 template_name: Optional[str], template_fields: dict, history: List[str]) -> str:
    history_part = ''
    if history:
        history_part = "\n---\nPrevious conversation:\n" + "\n".join(history[-10:]) + "\n---\n"
    if template_name and template_name in CORPORATE_TEMPLATES:
        template = CORPORATE_TEMPLATES[template_name].format(**template_fields)
        body = template.replace('{body}', user_text)
    else:
        body = user_text
    prompt_parts = [
        "You are a senior corporate communications assistant.",
        f"Input language: {input_lang}. Output language: {output_lang}.",
        f"Tone: {tone}.",
        "Produce a polished, professional message suitable for corporate use.",
        history_part,
        "Message:",
        body,
        "-- End of instructions --"
    ]
    return "\n".join([p for p in prompt_parts if p])

def generate_docx_from_text(text: str) -> bytes:
    doc = Document()
    for line in text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def make_eml(subject: str, body: str, sender: str='sender@example.com', recipient: str='recipient@example.com') -> bytes:
    eml = f"""From: {sender}
To: {recipient}
Subject: {subject}
MIME-Version: 1.0
Content-Type: text/plain; charset=utf-8

{body}"""
    return eml.encode('utf-8')

def serve_gtts_to_user(text: str) -> Optional[bytes]:
    if not gTTS:
        return None
    t = gTTS(text)
    bio = io.BytesIO()
    t.write_to_fp(bio)
    bio.seek(0)
    return bio.read()

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title=APP_TITLE, layout='centered')

if 'history' not in st.session_state:
    st.session_state.history = []

if 'last_generated' not in st.session_state:
    st.session_state.last_generated = ''

st.title(APP_TITLE)

input_lang_name = st.selectbox('Input language', list(LANGUAGES.keys()))
output_lang_name = st.selectbox('Output language', list(LANGUAGES.keys()))
tone = st.selectbox('Tone', TONE_OPTIONS)

uploaded = st.file_uploader('Upload source file', type=['txt','docx','pdf'])
uploaded_text = read_uploaded_file(uploaded) if uploaded else ''
user_text = st.text_area('Source text', value=uploaded_text, height=200)

template_choice = st.selectbox('Template', ['(none)'] + list(CORPORATE_TEMPLATES.keys()))
template_fields = {}
if template_choice != '(none)':
    template_fields['recipient'] = st.text_input('Recipient', 'Team')
    template_fields['sender'] = st.text_input('Sender', 'Your Name')
    template_fields['subject'] = st.text_input('Subject', 'Updates')
    template_fields['times'] = st.text_input('Proposed times', 'Mon/Tue 10-11am')

generate_button = st.button('Generate Message')
reset_button = st.button('Reset Memory')

if reset_button:
    st.session_state.history = []
    st.session_state.last_generated = ''
    st.success('Memory reset.')

if generate_button and genai_client:
    # Translate input to English for processing
    internal_input = user_text
    if input_lang_name != 'English' and _translator:
        internal_input = translate_text(user_text, 'en')

    prompt = build_prompt(internal_input, input_lang_name, output_lang_name, tone,
                          (template_choice if template_choice != '(none)' else None),
                          template_fields, st.session_state.history)

    # Generate with GenAI
    try:
        response = genai_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        result = response.candidates[0].content.parts[0].text
    except Exception as e:
        st.error(f"Generation failed: {e}")
        result = ''

    # Translate back to output language
    final_output = result
    if output_lang_name != 'English' and _translator:
        final_output = translate_text(result, LANGUAGES[output_lang_name])

    st.session_state.last_generated = final_output
    st.session_state.history.append(f"User request: {user_text}")
    st.session_state.history.append(f"Generated: {final_output}")
    st.success('Message generated.')

# Preview & downloads
if st.session_state.last_generated:
    escaped_text = st.session_state.last_generated.replace("`","\\`").replace("\n","\\n")
    tts_html = f"""
    <div>
    <button id='speak' style='background:#b8860b; border:none; padding:8px 12px; border-radius:6px; font-weight:700; color:#071018'>Speak</button>
    <script>
    const text = `{escaped_text}`;
    const btn = document.getElementById('speak');
    btn.onclick = () => {{
        const utter = new SpeechSynthesisUtterance(text);
        utter.lang='{LANGUAGES[output_lang_name]}';
        window.speechSynthesis.speak(utter);
    }};
    </script>
    </div>
    """
    st.markdown(tts_html, unsafe_allow_html=True)
    st.text_area('Generated Message', st.session_state.last_generated, height=200)
    st.download_button('Download TXT', st.session_state.last_generated, 'message.txt')
    st.download_button('Download DOCX', generate_docx_from_text(st.session_state.last_generated), 'message.docx')

st.sidebar.title('Session Memory')
if st.session_state.history:
    for i, item in enumerate(st.session_state.history[::-1]):
        st.sidebar.write(f"{len(st.session_state.history)-i}. {item[:200]}{'...' if len(item)>200 else ''}")
else:
    st.sidebar.write("No history yet.")
>>>>>>> 0f0227921e5af3e35664a0d860733d4ce4ff414e
